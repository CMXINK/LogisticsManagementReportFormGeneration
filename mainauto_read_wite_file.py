from collections import OrderedDict
from ctypes.wintypes import RGB
from datetime import datetime
# 读取数据
from docx.shared import Pt
from pyexcel_xls import save_data, get_data
from openpyxl import load_workbook
import os
import docx.document, docx


class Read_file():
    def excel(self, path):
        dic = OrderedDict()
        # 抓取数据
        xdata = get_data(path)
        for sheet in xdata:
            dic[sheet] = xdata[sheet]
        return dic


class Write_file():
    def excel(self, path, data):
        dic = OrderedDict()
        for sheetName, sheetValue in data.items():
            d = {sheetName: sheetValue}
            dic.update(d)

        save_data(path, dic)


class form_operation():

    def copy_template(self, template, data, rename=1):
        """

        :param template: 模板名称
        :param data: 需要插入表格的数据
        :param rename: 如果是一张表，需要传入修改的名称，应该带后缀，多张则传入张数
        :return: none
        """
        path = os.path.abspath(".") + "\\template\\" + template

        if type(rename) == str:
            save_path = os.path.abspath(".") + "\\Excel_files\\" + str(rename)
            content = 'copy {} {}'.format(path, save_path)
            os.system(content)
            print("path", path, "save_path=", save_path)
            self.insert_data(data, save_path)
        else:
            for i in range(rename):
                rename1 = template + datetime.now().strftime("%Y%M%d%H%M%S") + "_" + str(i) + ".xlsx"
                save_path = os.path.abspath(".") + "\\Excel_files\\" + rename1
                content = 'copy {} {}'.format(path, save_path)
                os.system(content)
                self.insert_data(data[i], save_path)

    def insert_data(self, data, save_path):
        """
        注意当前只对xlsx生效，如果是xls文件则需要另外调用（class Write_file()）并修改data数据  且  修改copy_template中的rename1的后缀
        """
        path = save_path
        wb = load_workbook(path)
        sheet = wb.active
        print(data)
        count = 0
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value:
                    cell.value = data[count]
                    count += 1
        wb.save(path)
        # os.remove(path)

    # 对于集装箱文件的操作
    def copy_jizhuangxiang_file(self, filename, save_name, data):
        save_name = save_name + '.xlsx'
        path = os.path.abspath(".") + "\\template\\" + "集装箱{}.xlsx".format(filename)
        save_path = os.path.abspath(".") + "\\Excel_files\\" + str(save_name)
        content = 'copy {} {}'.format(path, save_path)
        os.system(content)
        self.insert_jizhuangxiang_data(save_path, data)

    def insert_jizhuangxiang_data(self, save_path, data):
        """
        注意当前只对xlsx生效，如果是xls文件则需要另外调用（class Write_file()）并修改data数据  且  修改copy_template中的rename1的后缀
        """
        path = save_path
        wb = load_workbook(path)
        sheet = wb.active
        print(data)
        count = 0
        for row in sheet.iter_rows():
            for cell in row:
                try:
                    if int(cell.value) == 1:
                        cell.value = data[count]
                        count += 1
                except:
                    pass
        wb.save(path)

    # 对散货费用单的操作
    def copy_feiyongdan_file(self, filename, save_name, data):
        save_name = save_name + '.xlsx'
        path = os.path.abspath(".") + "\\template\\" + "费用{}.xlsx".format(filename)
        save_path = os.path.abspath(".") + "\\Excel_files\\" + str(save_name)
        content = 'copy {} {}'.format(path, save_path)
        os.system(content)
        self.insert_feiyongdan_data(save_path, data)

    def insert_feiyongdan_data(self, save_path, data):
        """
        注意当前只对xlsx生效，如果是xls文件则需要另外调用（class Write_file()）并修改data数据  且  修改copy_template中的rename1的后缀
        """
        path = save_path
        wb = load_workbook(path)
        sheet = wb.active
        print(data)
        count = 0
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value == "cmx":
                    try:
                        cell.value = data[count]
                        count += 1
                    except:
                        pass
        wb.save(path)

    def jizhuangxiang_docx_file(self, file_name, save_name, data):
        file_path = os.path.abspath(".") +"\\template\\" +"集装箱{}.docx".format(file_name)
        doc = docx.Document(file_path)
        check = 0
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if "c" in run.text:
                    run.text = run.text.replace('c', data[check])
                    check += 1

        doc.save(os.path.abspath(".") +"\\Excel_files\\" +save_name+".docx")


    def sanhuo_docx_file(self, file_name, save_name, data_par, date_box):
        file_path = os.path.abspath(".")+"\\template\\"+"{}下货纸.docx".format(file_name)
        # 对段落的操作
        doc = docx.Document(file_path)
        check = 0
        print(data_par)
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                try:
                    if "c" in run.text:
                        run.text = run.text.replace('c', data_par[check])
                        check += 1
                except:
                    pass

        file = doc
        children = file.element.body.iter()
        child_iters = []
        tags = []
        for child in children:
            # 通过类型判断目录
            if child.tag.endswith(('AlternateContent', 'textbox')):
                for ci in child.iter():
                    tags.append(ci.tag)
                    if ci.tag.endswith(('main}r', 'main}pPr')):
                        child_iters.append(ci)
        text = ['']
        for ci in child_iters:
            if ci.tag.endswith('main}pPr'):
                text.append('')
            else:
                text[-1] += ci.text
            ci.text = ''

        i, k = 0, 0
        count = 0
        for ci in child_iters:
            if ci.tag.endswith('main}pPr'):
                i += 1
                k = 0
            elif k == 0:
                ci.text = date_box[count]
                count += 1
                k = 1
        doc.save(os.path.abspath(".") + "\\Excel_files\\" + save_name+".docx")
        # 对文本框的操作

if __name__ == '__main__':
    pass
